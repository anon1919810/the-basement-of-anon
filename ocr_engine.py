# -*- coding: utf-8 -*-
"""
文本/OCR引擎模块：PDF文本层、Word、扫描件OCR（RapidOCR/Tesseract）、结构化标题
"""
import os

import pymupdf as fitz

import config
from shared import _ZIMU_RE


def _clean_ocr_text(text):
    """过滤OCR噪音行：丢弃纯乱码/边注行（非中文占比过高的行），保留正文"""
    lines = []
    for line in text.splitlines():
        s = line.strip()
        if not s:
            continue
        cjk = sum(1 for ch in s if '\u4e00' <= ch <= '\u9fff')
        # 至少2个汉字，或含1个汉字且中文占比>=30%才保留
        if cjk >= 2 or (cjk >= 1 and cjk / len(s) >= 0.3):
            lines.append(s)
    return "\n".join(lines)


def ocr_pdf(pdf_path):
    """使用OCR识别扫描版PDF（优先RapidOCR：中文准确率远高于Tesseract，且能识别【】标记；
    失败自动回退Tesseract。灰度化 + 分批转换防内存溢出 + 噪音过滤 + 进度输出）"""
    try:
        from pdf2image import convert_from_path
        import pytesseract

        # 仅当配置了本地存在的Tesseract路径时覆盖；否则交给PATH查找（云端由packages.txt安装）
        if config.TESSERACT_PATH and os.path.exists(config.TESSERACT_PATH):
            pytesseract.pytesseract.tesseract_cmd = config.TESSERACT_PATH

        # 惰性加载 RapidOCR 引擎（首次调用初始化，约1-2秒；未安装则自动回退Tesseract）
        rapid_engine = [None]

        def _rapid(img):
            if rapid_engine[0] is None:
                from rapidocr_onnxruntime import RapidOCR
                rapid_engine[0] = RapidOCR()
            import numpy as np
            result, _ = rapid_engine[0](np.array(img.convert("RGB")))
            if not result:
                return ""
            return "\n".join(line[1] for line in result)

        with fitz.open(pdf_path) as doc:
            n_pages = doc.page_count

        full_text = ""
        batch = 10  # 每批最多10页，防止大PDF一次性载入过多图片
        for start in range(1, n_pages + 1, batch):
            end = min(start + batch - 1, n_pages)
            images = convert_from_path(pdf_path, dpi=config.OCR_DPI,
                                       first_page=start, last_page=end)
            for img in images:
                gray = img.convert("L")  # 灰度化提升识别率
                text = ""
                if config.OCR_ENGINE in ("auto", "rapidocr"):
                    try:
                        text = _rapid(gray)
                    except Exception as e:
                        print(f"  [OCR] RapidOCR不可用({type(e).__name__})，回退Tesseract", flush=True)
                if not text and config.OCR_ENGINE in ("auto", "tesseract"):
                    text = pytesseract.image_to_string(gray, lang='chi_sim+eng')
                full_text += text + "\n"
            print(f"  [OCR] 进度 {end}/{n_pages} 页", flush=True)
        return _clean_ocr_text(full_text)
    except ImportError:
        print("[警告] 未安装pdf2image或pytesseract，OCR功能不可用")
        return ""
    except Exception as e:
        print(f"[警告] OCR识别失败：{e}")
        return ""


def extract_text_from_pdf(pdf_path):
    """从PDF提取文本（对外接口，返回文本字符串）"""
    text, _ = _extract_text_with_flag(pdf_path)
    return text


def _extract_docx(file_path):
    """从Word(.docx)提取文本：直接读取段落，完全跳过OCR（效果等同文本层PDF）"""
    try:
        from docx import Document
        doc = Document(file_path)
        parts = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)
        # 表格文本（志书常见"附表"）
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append("｜".join(cells))
        return "\n\n".join(parts)
    except ImportError:
        print("[警告] 未安装python-docx，无法读取Word文件")
        return ""
    except Exception as e:
        print(f"[警告] Word读取失败：{e}")
        return ""


def _extract_structured_text(pdf_path):
    """字体感知提取：把"加粗/大字提行标题"识别为子目并插入【】标记。

    用于文本层PDF无【】但存在加粗/大字标题的情况（如"亭台楼阁""古建筑"类提行标题）。
    返回带【】标记的结构化文本；无标题结构时返回空（调用方回退原文）。
    """
    try:
        doc = fitz.open(pdf_path)
        # 第一遍：统计字号，取中位数作为正文字号
        sizes = []
        for page in doc:
            for b in page.get_text("dict").get("blocks", []):
                if b.get("type") != 0:
                    continue
                for line in b.get("lines", []):
                    for span in line.get("spans", []):
                        if span.get("text", "").strip():
                            sizes.append(span.get("size", 0))
        doc.close()
        if not sizes:
            return ""
        body_size = sorted(sizes)[len(sizes) // 2]
        threshold = max(body_size * 1.35, body_size + 3.0)

        # 第二遍：重建文本，标题行加【】包裹
        doc = fitz.open(pdf_path)
        out_lines = []
        for page in doc:
            for b in page.get_text("dict").get("blocks", []):
                if b.get("type") != 0:
                    continue
                for line in b.get("lines", []):
                    spans = line.get("spans", [])
                    if not spans:
                        continue
                    stripped = "".join(s.get("text", "") for s in spans).strip()
                    if not stripped:
                        continue
                    s0 = spans[0]
                    bold = bool(s0.get("flags", 0) & 16)  # bit4=加粗
                    size = s0.get("size", 0)
                    is_title = (bold or size >= threshold) and len(stripped) <= 24
                    if is_title:
                        out_lines.append(f"【{stripped}】")
                    else:
                        out_lines.append(stripped)
            out_lines.append("")
        doc.close()
        return "\n".join(out_lines)
    except Exception as e:
        print(f"  [警告] 结构化提取失败：{e}")
        return ""


def _extract_text_with_flag(file_path):
    """提取文本，返回 (text, used_ocr)。支持PDF与Word(.docx)"""
    if file_path.lower().endswith(".docx"):
        return _extract_docx(file_path), False

    used_ocr = False
    try:
        doc = fitz.open(file_path)
        full_text = ""
        for page in doc:
            # sort=True 按阅读顺序排序，改善多栏/表格版面
            page_text = page.get_text("text", sort=True)
            full_text += page_text + "\n\n"
        doc.close()

        # 文本量极少但PDF有内容，启用OCR
        if len(full_text.strip()) < 100 and config.ENABLE_OCR:
            print(f"  [提示] {os.path.basename(file_path)} 疑似扫描件，启动OCR...")
            full_text = ocr_pdf(file_path)
            used_ocr = True
        elif len(_ZIMU_RE.findall(full_text)) < 3:
            # 文本层无【】结构 -> 尝试字体感知的加粗/大字标题结构
            structured = _extract_structured_text(file_path)
            if len(_ZIMU_RE.findall(structured)) >= 3:
                print(f"  [结构] 检测到加粗/大字标题子目，使用结构化提取")
                full_text = structured

        return full_text, used_ocr
    except Exception as e:
        print(f"  [警告] 读取PDF失败：{e}")
        return "", False


__all__ = [
    "_clean_ocr_text", "ocr_pdf", "extract_text_from_pdf",
    "_extract_docx", "_extract_structured_text", "_extract_text_with_flag",
]
